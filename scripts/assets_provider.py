# -*- coding: utf-8 -*-
"""
assets_provider.py · B 模式 assets 库对接抽象层(V61 基建)

定义 AssetsProvider 抽象接口 + PlaceholderAssetsProvider 本轮实现。

B 模式 b_mode_fill 通过本接口查找和解析 assets 条目,不直接操作 assets 目录,
为未来材料管理子系统(α 路线,见 business_model §7.4)预留接口位置。

【未来实现(当前不写,仅接口注释说明)】

CuratedLocalAssetsProvider:
    对接本地 assets/ 目录,由材料管理子系统(assets_curator.py)预先归排。
    lookup 按 asset_type + 项目上下文语义查找;resolve 返回真实 docx/pdf 路径。
    触发条件:B 模式在真实项目中稳定运行 + 材料管理子系统上线。

MCPExternalAssetsProvider:
    通过 MCP 协议对接外部材料库(如企业统一材料平台)。
    适用于大型组织,资质/业绩材料在统一平台集中管理。
    触发条件:有真实外部材料库需求时。
"""

from __future__ import annotations

from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
import tempfile


@dataclass
class AssetRef:
    """Assets 库中一个资源的引用。"""
    asset_type: str              # 资源类型,如 "资质证明" / "类似业绩" / "人员简历"
    is_placeholder: bool         # True = 占位引用,False = 真实命中
    lookup_key: str              # 查询键,便于未来查询溯源
    metadata: dict = field(default_factory=dict)  # 自由字段,供未来实现扩展


class AssetsProvider(ABC):
    """B 模式 assets 库对接抽象接口。"""

    @abstractmethod
    def lookup(self, asset_type: str, **kwargs) -> AssetRef:
        """按 asset_type 和可选参数查找资源,返回 AssetRef。"""
        ...

    @abstractmethod
    def resolve(self, asset_ref: AssetRef) -> Path:
        """把 AssetRef 解析为实际可用文件路径(docx / pdf 等)。"""
        ...


class PlaceholderAssetsProvider(AssetsProvider):
    """V61 本轮实现:所有 lookup 返回占位 AssetRef,resolve 产出占位 docx。

    不查任何真实资源。用于 B 模式基建验证和材料管理子系统启动前的占位运行。
    """

    def lookup(self, asset_type: str, **kwargs) -> AssetRef:
        # 构造查询键(便于溯源,即使本轮不用)
        kv_parts = [f"{k}={v}" for k, v in sorted(kwargs.items())]
        lookup_key = f"{asset_type}|{'|'.join(kv_parts)}" if kv_parts else asset_type
        return AssetRef(
            asset_type=asset_type,
            is_placeholder=True,
            lookup_key=lookup_key,
            metadata={"reason": "PlaceholderAssetsProvider: 本轮不做真实查找"},
        )

    def resolve(self, asset_ref: AssetRef) -> Path:
        """生成含占位文字的临时 docx 段落文件,存放于临时目录(不进 tracked_outputs)。"""
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("缺少 python-docx 依赖")

        doc = Document()
        doc.add_paragraph(f"[此处插入 {asset_ref.asset_type} 材料]")
        doc.add_paragraph(f"  lookup_key: {asset_ref.lookup_key}")
        doc.add_paragraph("  (由 PlaceholderAssetsProvider 产出,真实材料待填)")

        # 临时目录,不进基线
        tmp_dir = Path(tempfile.gettempdir()) / "tender_writer_placeholder_assets"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        # 用 lookup_key 的哈希做文件名避免冲突
        import hashlib
        fname = f"placeholder_{hashlib.md5(asset_ref.lookup_key.encode('utf-8')).hexdigest()[:12]}.docx"
        out_path = tmp_dir / fname
        doc.save(str(out_path))
        return out_path


def get_provider(name: str) -> AssetsProvider:
    """按 manifest.yaml 的 assets_provider 字段返回对应实例。"""
    if name == "placeholder":
        return PlaceholderAssetsProvider()
    raise ValueError(
        f"未知的 assets_provider='{name}'。"
        f"当前仅支持 'placeholder';未来的 CuratedLocalAssetsProvider / "
        f"MCPExternalAssetsProvider 见本模块文档。"
    )
