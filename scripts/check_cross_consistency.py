# -*- coding: utf-8 -*-
"""
check_cross_consistency.py · 跨章节一致性检查

用途:
    在 compliance_check.py 之前跑一次,捕获 AI 写作中常见的"自相矛盾"类错误。
    典型案例:团队规模 17 人 × 人天成本 → 超预算;承诺 D+65 超 60 天工期。

与 compliance_check 的关系:
    compliance_check 关注"评分项是否覆盖 / ★▲ 是否响应 / 模板残留 / 格式",
    本脚本关注"AI 输出的数字/时间/金额之间是否自相矛盾"。两者互补,都要跑。

用法:
    ./run_script.bat check_cross_consistency.py \\
        projects/{项目}/output/tender_response.docx \\
        --brief projects/{项目}/output/tender_brief.json \\
        --per-person-day-yuan 1500  # 可选,默认 1500 元/人·天(行业中位数)
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from docx import Document


# ─────────────────────────────────────────────
# 数字抽取工具
# ─────────────────────────────────────────────

_MONEY_RE = re.compile(r"(\d+(?:\.\d+)?)\s*万元|(\d{4,})\s*元")
_DAYS_RE = re.compile(r"(\d+)\s*(日历日|天|个日历日)")
_PERSON_RE = re.compile(r"(\d+)\s*(人|位)")
_D_PLUS_RE = re.compile(r"D\s*\+\s*(\d+)")

# v2.0.1:团队语境词表。只有匹配"N 人/N 位"前 20 字窗口里含以下任一词,
# 才视为团队规模数字,纳入成本估算。否则视为培训规模 / 服务对象 / 无关计数,过滤掉。
_TEAM_CONTEXT_WORDS = (
    "项目组", "团队", "项目经理", "技术骨干", "核心人员",
    "专家", "工程师", "成员", "人员配置", "项目部",
)


def extract_docx_text(docx_path: Path) -> str:
    """把 docx 的段落和表格单元格拼成纯文本,供后续正则抽取。"""
    doc = Document(str(docx_path))
    chunks = []
    for p in doc.paragraphs:
        chunks.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def parse_money_in_yuan(text: str) -> list[tuple[float, str]]:
    """返回 (金额元, 原文片段) 列表。万元 × 10000。"""
    out = []
    for m in _MONEY_RE.finditer(text):
        wan = m.group(1)
        yuan = m.group(2)
        if wan:
            try:
                out.append((float(wan) * 10000, m.group(0)))
            except ValueError:
                pass
        elif yuan:
            try:
                out.append((float(yuan), m.group(0)))
            except ValueError:
                pass
    return out


def parse_days(text: str) -> list[tuple[int, str]]:
    return [(int(m.group(1)), m.group(0)) for m in _DAYS_RE.finditer(text)]


def parse_d_plus(text: str) -> list[int]:
    return [int(m.group(1)) for m in _D_PLUS_RE.finditer(text)]


def parse_people(text: str) -> list[int]:
    """
    v2.0.1:只保留"N 人/N 位"前 20 字窗口里含团队语境词的匹配。
    过滤培训规模(如"培训不少于 100 人")/ 服务对象(如"服务 3000 万群众")/
    其他无关计数,避免团队成本估算误判。
    """
    out = []
    for m in _PERSON_RE.finditer(text):
        start = max(0, m.start() - 20)
        ctx = text[start:m.start()]
        if any(w in ctx for w in _TEAM_CONTEXT_WORDS):
            out.append(int(m.group(1)))
    return out


# ─────────────────────────────────────────────
# 检查项
# ─────────────────────────────────────────────

def check_duration_vs_dplus(budget_days: int | None,
                            d_plus_values: list[int]) -> list[str]:
    """
    检查承诺时间节点(D+N)不得超过服务期限(budget_days)。

    分级阈值:
    - D+N > duration_days × 2 → [失败](明显矛盾,如 D+130 对 60 天工期)
    - duration_days < D+N ≤ duration_days × 2 → [警告](可能是后续服务/评审配合,需人工复核)
    - D+N ≤ duration_days → [通过]
    """
    issues = []
    if budget_days is None or budget_days <= 0:
        issues.append(f"[信息] 未能从 tender_brief 获取服务期限,跳过 D+N vs 工期检查")
        return issues
    uniq = sorted(set(d_plus_values))
    hard_violations = [d for d in uniq if d > budget_days * 2]
    soft_violations = [d for d in uniq
                       if budget_days < d <= budget_days * 2]
    for d in hard_violations:
        issues.append(
            f"[失败] D+{d} 明显超出服务期限 {budget_days} 天(超 {d - budget_days} 天),"
            f"明显矛盾"
        )
    for d in soft_violations:
        issues.append(
            f"[警告] D+{d} 超出服务期限 {budget_days} 天(超 {d - budget_days} 天),"
            f"若属后续服务/评审配合可忽略,属承诺节点需修改"
        )
    if not hard_violations and not soft_violations:
        issues.append(
            f"[通过] docx 中 D+N 节点({uniq})均在服务期 {budget_days} 天内"
        )
    return issues


def check_team_cost_vs_budget(people_counts: list[int],
                              duration_days: int | None,
                              budget_yuan: float | None,
                              per_person_day_yuan: float = 1500.0,
                              tolerance: float = 1.5) -> list[str]:
    """
    检查团队规模 × 人天成本 是否在合同预算 × tolerance 内。

    计算:总人天 = 团队总人数 × 工期天数(取 docx 中最大团队规模 + 预算工期)
          估算成本 = 总人天 × per_person_day_yuan
    告警阈值:估算成本 > 合同预算 × tolerance(默认 1.5 倍)
    """
    issues = []
    if not people_counts:
        issues.append("[信息] docx 中未抽到团队人数,跳过团队成本检查")
        return issues
    if duration_days is None or duration_days <= 0:
        issues.append("[信息] 未获取服务期限,跳过团队成本检查")
        return issues
    if budget_yuan is None or budget_yuan <= 0:
        issues.append("[信息] 未获取合同预算,跳过团队成本检查")
        return issues

    max_people = max(people_counts)
    # 假设团队 50% 投入,考虑到有专家只是部分时间参与
    effective_people = max_people * 0.5
    estimated_cost = effective_people * duration_days * per_person_day_yuan
    threshold = budget_yuan * tolerance

    status = "失败" if estimated_cost > threshold else "通过"
    issues.append(
        f"[{status}] 团队成本估算: "
        f"{max_people} 人 × 50% 投入率 × {duration_days} 天 × {per_person_day_yuan} 元/人·天 "
        f"= {estimated_cost:,.0f} 元;合同预算 {budget_yuan:,.0f} 元(含 {tolerance}× 容忍度)"
    )
    if estimated_cost > threshold:
        issues.append(
            f"[建议] 团队规模/投入率/工期/预算至少一个维度需调整;"
            f"人数最大值 {max_people} 可能偏大,检查是否含专家或外协"
        )
    return issues


def check_key_numbers_consistency(text: str, budget_yuan: float | None) -> list[str]:
    """
    检查 docx 中出现的关键数字与 tender_brief 是否一致。
    若 docx 出现的金额 > 预算 10 倍,判失败(明显矛盾);1-10 倍内判警告(可能是分项/保险额等)。
    """
    issues = []
    if budget_yuan is None or budget_yuan <= 0:
        return issues
    docx_moneys = parse_money_in_yuan(text)
    hard = [m for m, _ in docx_moneys if m > budget_yuan * 10]
    soft = [m for m, _ in docx_moneys
            if budget_yuan < m <= budget_yuan * 10]
    if hard:
        uniq = sorted(set(hard))[:5]
        issues.append(
            f"[失败] docx 中金额 {uniq} 大幅超出预算 {budget_yuan:,.0f} 元(10 倍以上),"
            f"疑似编造或单位错误"
        )
    if soft:
        uniq = sorted(set(soft))[:5]
        issues.append(
            f"[警告] docx 中金额 {uniq} 超出预算 {budget_yuan:,.0f} 元(1-10 倍),"
            f"若属分项/保险额/行业参考金额可忽略"
        )
    return issues


# ─────────────────────────────────────────────
# 主入口
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="跨章节一致性检查(v2)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("docx_path", help="tender_response.docx 或 final_response.docx 路径")
    parser.add_argument("--brief", required=True, help="tender_brief.json 路径")
    parser.add_argument("--per-person-day-yuan", type=float, default=1500.0,
                        help="人天成本估算基准(元/人·天,默认 1500)")
    parser.add_argument("--tolerance", type=float, default=1.5,
                        help="成本告警容忍度倍数(默认 1.5,即估算 > 预算 × 1.5 告警)")
    args = parser.parse_args()

    docx_path = Path(args.docx_path)
    brief_path = Path(args.brief)

    if not docx_path.exists():
        print(f"[错误] 找不到 docx: {docx_path}", file=sys.stderr)
        sys.exit(1)
    if not brief_path.exists():
        print(f"[错误] 找不到 tender_brief.json: {brief_path}", file=sys.stderr)
        sys.exit(1)

    # v2 单元 7 P16:用函数级闸门读 brief(而不是直接 json.loads 绕过 gate)
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from brief_schema import load_brief_guarded
    brief_data = load_brief_guarded(brief_path)
    extracted = brief_data.get("extracted", {})

    # 解析服务期限天数(从 extracted.duration 抽)
    duration_str = extracted.get("duration", "") or ""
    duration_match = _DAYS_RE.search(duration_str)
    duration_days = int(duration_match.group(1)) if duration_match else None

    # 解析合同预算金额(从 extracted.budget 抽)
    # budget_str 可能形如:"500000.00元(人民币伍拾万元整)" 或 "50万元" 或 "50 万元"
    # 规则:第一个带单位的数字决定金额;单位=元直接用,单位=万元/万 乘以 10000
    budget_str = (extracted.get("budget", "") or "").replace(",", "")
    budget_yuan = None
    m_wanyuan = re.search(r"(\d+(?:\.\d+)?)\s*万元", budget_str)
    m_yuan = re.search(r"(\d+(?:\.\d+)?)\s*元", budget_str)
    if m_wanyuan and (not m_yuan or m_wanyuan.start() < m_yuan.start()):
        budget_yuan = float(m_wanyuan.group(1)) * 10000
    elif m_yuan:
        budget_yuan = float(m_yuan.group(1))

    # 抽 docx 文本
    text = extract_docx_text(docx_path)

    d_plus_values = parse_d_plus(text)
    people_counts = parse_people(text)

    print("=" * 60)
    print("跨章节一致性检查报告")
    print("=" * 60)
    print(f"docx:  {docx_path}")
    print(f"brief: {brief_path}")
    print(f"服务期限: {duration_days} 天" if duration_days else "服务期限: 未识别")
    print(f"合同预算: {budget_yuan:,.0f} 元" if budget_yuan else "合同预算: 未识别")
    print(f"docx 抽到 D+N 节点: {sorted(set(d_plus_values))}")
    print(f"docx 抽到团队人数样本: {sorted(set(people_counts))[:10]}")
    print()

    all_issues = []

    print("[检查 1] 承诺时间节点 vs 服务期限")
    issues = check_duration_vs_dplus(duration_days, d_plus_values)
    for msg in issues:
        print(f"  {msg}")
    all_issues.extend(issues)
    print()

    print("[检查 2] 团队成本估算 vs 合同预算")
    issues = check_team_cost_vs_budget(
        people_counts, duration_days, budget_yuan,
        per_person_day_yuan=args.per_person_day_yuan,
        tolerance=args.tolerance,
    )
    for msg in issues:
        print(f"  {msg}")
    all_issues.extend(issues)
    print()

    print("[检查 3] docx 中的金额数字与预算一致性")
    issues = check_key_numbers_consistency(text, budget_yuan)
    if not issues:
        print("  [通过] 未发现明显异常金额")
    for msg in issues:
        print(f"  {msg}")
    all_issues.extend(issues)
    print()

    # 汇总
    fail_count = sum(1 for m in all_issues if m.startswith("[失败]"))
    warn_count = sum(1 for m in all_issues if m.startswith("[警告]"))
    print("=" * 60)
    print(f"跨章节一致性检查汇总: {fail_count} 失败 / {warn_count} 警告")
    print("=" * 60)

    if fail_count > 0:
        sys.exit(1)


if __name__ == "__main__":
    main()
